import sqlite3
import matplotlib.pyplot as plt
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import Inches
import os
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from ibm_watsonx_ai import Credentials
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from ibm_watsonx_ai.foundation_models.utils.enums import DecodingMethods, ModelTypes
from langchain_ibm import WatsonxLLM

class CarbonFootprintReport:
    def __init__(self, db_path='ocr_results.db', doc_path='report_carbon_footprint_user.docx'):
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.doc = Document()  # Initialize a Word document
        self.doc_path = doc_path
        self.watsonx_llm = self.setup_watsonx()

    def setup_watsonx(self):
        try:
            print("Setting up WatsonX...")
            credentials = Credentials(
                url="https://us-south.ml.cloud.ibm.com",
                api_key="add your api key here"  # Replace with your actual API key
            )
            project_id = os.environ.get("PROJECT_ID", "add your project id here")  # Set your project ID here
            model_id = ModelTypes.GRANITE_13B_CHAT_V2
            parameters = {
                GenParams.DECODING_METHOD: DecodingMethods.GREEDY,
                GenParams.MIN_NEW_TOKENS: 1,
                GenParams.MAX_NEW_TOKENS: 500,
                GenParams.STOP_SEQUENCES: ["Human:", "Receipt:", "Total"]
            }
            watsonx_llm = WatsonxLLM(
                model_id=model_id.value,
                url=credentials.get("url"),
                apikey=credentials.get("apikey"),
                project_id=project_id,
                params=parameters
            )
            print("WatsonX set up successfully.")
            return watsonx_llm
        except Exception as e:
            print(f"Error setting up WatsonX: {str(e)}")
            return None

    def save_pie_chart(self, categories, carbon_values, filename):
        # Plot Pie Chart and save as an image
        plt.figure(figsize=(8, 6))
        plt.pie(carbon_values, labels=categories, autopct='%1.1f%%', startangle=140)
        plt.title("Category-wise Carbon Footprint Distribution")
        plt.axis('equal')
        plt.savefig(filename)
        plt.close()

    def save_bar_chart(self, dates, carbon_values, filename):
        # Plot Bar Graph and save as an image
        plt.figure(figsize=(10, 6))
        plt.bar(dates, carbon_values, color='skyblue')
        plt.xticks(rotation=45, ha='right')
        plt.title("Date-wise Carbon Footprint")
        plt.xlabel("Date")
        plt.ylabel("Carbon Footprint")
        plt.tight_layout()
        plt.savefig(filename)
        plt.close()

    def generate_category_wise_report(self):
        self.cursor.execute('''SELECT category, SUM(carbon_footprint) as total_carbon FROM receipts GROUP BY category''')
        result = self.cursor.fetchall()

        # Display textual report
        self.doc.add_heading("Category-wise Carbon Footprint Report", level=1)
        categories = []
        carbon_values = []
        for row in result:
            categories.append(row[0])
            carbon_values.append(row[1])
            self.doc.add_paragraph(f"Category: {row[0]}, Total Carbon Footprint: {row[1]:.2f}")

        # Save Pie Chart to image and insert into document
        pie_chart_image = 'category_pie_chart.png'
        self.save_pie_chart(categories, carbon_values, pie_chart_image)
        self.doc.add_picture(pie_chart_image, width=Inches(5))

    def generate_date_wise_report(self):
        self.cursor.execute('''SELECT current_date, SUM(carbon_footprint) as total_carbon FROM receipts GROUP BY current_date''')
        result = self.cursor.fetchall()

        # Display textual report
        self.doc.add_heading("Date-wise Carbon Footprint Report", level=1)
        dates = []
        carbon_values = []
        for row in result:
            date = datetime.strptime(row[0], '%Y-%m-%d').strftime('%B %d, %Y')
            dates.append(date)
            carbon_values.append(row[1])
            self.doc.add_paragraph(f"Date: {date}, Total Carbon Footprint: {row[1]:.2f}")

        # Save Bar Graph to image and insert into document
        bar_chart_image = 'date_bar_chart.png'
        self.save_bar_chart(dates, carbon_values, bar_chart_image)
        self.doc.add_picture(bar_chart_image, width=Inches(5))

    def generate_combined_report(self):
        self.cursor.execute('''SELECT current_date, category, SUM(carbon_footprint) as total_carbon FROM receipts GROUP BY current_date, category''')
        result = self.cursor.fetchall()

        # Display textual report
        self.doc.add_heading("Combined Date-wise and Category-wise Report", level=1)
        for row in result:
            date = datetime.strptime(row[0], '%Y-%m-%d').strftime('%B %d, %Y')
            self.doc.add_paragraph(f"Date: {date}, Category: {row[1]}, Total Carbon Footprint: {row[2]:.2f}")

    def generate_detailed_report(self):
        self.cursor.execute('''SELECT item_name, carbon_footprint, category, current_date FROM receipts''')
        result = self.cursor.fetchall()

        # Display textual report
        self.doc.add_heading("Detailed Carbon Footprint Report", level=1)
        for row in result:
            date = datetime.strptime(row[3], '%Y-%m-%d').strftime('%B %d, %Y')
            self.doc.add_paragraph(f"Item: {row[0]}, Carbon Footprint: {row[1]:.2f}, Category: {row[2]}, Date: {date}")

    def generate_personalized_suggestions(self, report_summary):
        prompt_template = PromptTemplate(
            input_variables=["report"],
            template="""
            Based on the following carbon footprint report, provide personalized suggestions
            for reducing carbon footprint. The report contains details of items purchased,
            their carbon footprints, and their categories.

            Report:
            {report}

            Suggestions:
            """
        )

        llm_chain = LLMChain(llm=self.watsonx_llm, prompt=prompt_template)
        suggestions = llm_chain.run(report=report_summary)
        return suggestions

    def generate_complete_report(self):
        # Generate different sections of the report
        self.generate_category_wise_report()
        self.generate_date_wise_report()
        self.generate_combined_report()
        self.generate_detailed_report()

        # Create a summary for LLM
        summary_lines = []
        for line in self.doc.paragraphs:
            summary_lines.append(line.text)
        report_summary = "\n".join(summary_lines)

        # Get personalized suggestions from the LLM
        suggestions = self.generate_personalized_suggestions(report_summary)

        # Add suggestions to the report
        self.doc.add_heading("Personalized Suggestions", level=1)
        self.doc.add_paragraph(suggestions)

        # Save the report to a Word document
        self.save_report()

        # Clean up (optional: remove the chart images after saving the report)
        if os.path.exists('category_pie_chart.png'):
            os.remove('category_pie_chart.png')
        if os.path.exists('date_bar_chart.png'):
            os.remove('date_bar_chart.png')

    def save_report(self):
        # Save the Word document
        self.doc.save(self.doc_path)
        print(f"Report saved to {self.doc_path}")

    def close(self):
        self.conn.close()

    @classmethod
    def run_report(cls):
        # Path to your SQLite database and document
        db_path = 'ocr_results.db'
        doc_path = 'report_carbon_footprint_user.docx'

        # Initialize the report
        report = cls(db_path, doc_path)

        # Generate the complete report
        report.generate_complete_report()

        # Close the database connection
        report.close()

# Run the report if this script is executed
if __name__ == '__main__':
    CarbonFootprintReport.run_report()
