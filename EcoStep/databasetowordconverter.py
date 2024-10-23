import sqlite3
from docx import Document
import os
from datetime import datetime

# Path to your SQLite database
db_path = 'ocr_results.db'

class Database:
    def __init__(self, db_path):
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.cursor = self.conn.cursor()

    def fetch_all_receipts(self):
        """Fetch all records from the receipts table."""
        try:
            self.cursor.execute("SELECT * FROM receipts")
            return self.cursor.fetchall()
        except sqlite3.Error as e:
            print(f"An error occurred: {e}")
            return []

    def close(self):
        self.conn.close()

def export_receipt_data_to_word(receipts_data, output_path):
    """Appends receipt data along with the current date to the existing Word document."""
    
    # Load the existing Word document or create a new one if it doesn't exist
    if os.path.exists(output_path):
        document = Document(output_path)
    else:
        document = Document()
        document.add_heading('Receipt Information', level=1)
    
    # Get the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    
    # Iterate through each receipt data entry and add it to the document
    for receipt in receipts_data:
        item_name = receipt.get('item_name')
        carbon_footprint = receipt.get('carbon_footprint')
        category = receipt.get('category')
        
        # Add the current date before each receipt
        document.add_paragraph(f"Date: {current_date}")
        
        # Add receipt details
        document.add_paragraph(f"Item Name: {item_name}")
        document.add_paragraph(f"Carbon Footprint: {carbon_footprint} kg CO2e")
        document.add_paragraph(f"Category: {category}")
        document.add_paragraph("\n")
    
    # Save the updated document
    document.save(output_path)
    print(f"Data appended to {output_path}")
